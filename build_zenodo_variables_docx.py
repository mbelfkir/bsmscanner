from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "/Users/mbelfkir/HEP/BSMScanner/HHyybb_Zenodo_Dataset_Variables.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, col_widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(col_widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in col_widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, col_widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.10):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def format_run(run, size=11, bold=False, color="000000"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=16 if level == 1 else 12, after=8 if level == 1 else 6)
    r = p.add_run(text)
    format_run(r, size=16 if level == 1 else 13, bold=True, color="2E74B5")
    return p


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=6)
    r = p.add_run(text)
    format_run(r)
    return p


def write_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line=1.05)
    r = p.add_run(text)
    format_run(r, size=9.5, bold=bold)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
styles["Normal"].font.size = Pt(11)

title = doc.add_paragraph()
set_paragraph_spacing(title, after=8)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title_run = title.add_run("HH to bb gamma gamma ML Datasets")
format_run(title_run, size=18, bold=True, color="1F4D78")

subtitle = doc.add_paragraph()
set_paragraph_spacing(subtitle, after=12)
subtitle_run = subtitle.add_run("Zenodo dataset description and variable definitions")
format_run(subtitle_run, size=11, color="555555")

add_heading(doc, "Dataset Description", 1)
add_body_paragraph(
    doc,
    "This dataset contains simulated event samples for machine-learning studies of Higgs boson pair production in the HH -> bb gamma gamma final state. "
    "The samples include Standard Model gluon-fusion and vector-boson-fusion HH production, anomalous VBF HH scenarios with varied kappa_2V, and the main background processes: gamma gamma + jets, single-Higgs production, ttH, and ZH.",
)
add_body_paragraph(
    doc,
    "Each HDF5 file stores an event-level table with reconstructed object kinematics, event weights, topology variables, boosted/resolved category flags, and reconstructed Higgs-pair observables. "
    "The dataset is intended for training and benchmarking classifiers such as XGBoost, graph neural networks, transformer-based graph models, and image-based voxel models.",
)

add_heading(doc, "Variable Definitions", 1)

rows = [
    ("Photons", "`photon1_pt`, `photon1_eta`, `photon1_phi`, `photon2_pt`, `photon2_eta`, `photon2_phi`", "Transverse momentum, pseudorapidity, and azimuthal angle of the leading and sub-leading photons."),
    ("Resolved b-jets", "`bjet1_pt`, `bjet1_eta`, `bjet1_phi`, `bjet1_m`, `bjet2_pt`, `bjet2_eta`, `bjet2_phi`, `bjet2_m`", "Kinematics and mass of the leading and sub-leading small-radius b-tagged jets."),
    ("Boosted large-R jet", "`fatjet1_pt`, `fatjet1_eta`, `fatjet1_phi`, `fatjet1_m`, `fatjet1_flavor`, `fatjet1_btag`", "Kinematics, flavor label, and b-tagging information for the leading large-radius jet used in boosted selections."),
    ("Resolved VBF jets", "`VBFjet1_pt`, `VBFjet1_eta`, `VBFjet1_phi`, `VBFjet1_m`, `VBFjet2_pt`, `VBFjet2_eta`, `VBFjet2_phi`, `VBFjet2_m`", "Kinematics and mass of the two VBF-tagging jets in the resolved topology."),
    ("Boosted VBF jets", "`boostedVBFjet1_*`, `boostedVBFjet2_*`", "Kinematics and mass of the VBF-tagging jets used in the boosted topology."),
    ("VBF topology", "`mjj`, `deta`, `boostedmjj`, `boosteddeta`", "Dijet invariant mass and absolute pseudorapidity separation of the VBF jet pair, for resolved and boosted selections."),
    ("Missing transverse momentum", "`met_et`, `met_phi`", "Magnitude and azimuthal angle of missing transverse momentum."),
    ("Event weights", "`event_weight`, `event_lumi`, `event_xsec`, `event_br`", "Per-event weight, luminosity normalization, process cross section, and branching-ratio factor."),
    ("Event counts", "`event_ncentral`, `event_nforward`, `event_ntrks`, `event_nfatjets`", "Number of central jets, forward jets, tracks, and large-radius jets in the event."),
    ("Event categories", "`event_isBoosted`, `event_isResolved`", "Boolean flags identifying whether the event passes the boosted or resolved topology selection."),
    ("Diphoton Higgs candidate", "`Hyy_pt`, `Hyy_eta`, `Hyy_phi`, `Hyy_m`", "Reconstructed four-vector components of the H -> gamma gamma candidate."),
    ("bb Higgs candidate", "`Hbb_pt`, `Hbb_eta`, `Hbb_phi`, `Hbb_m`", "Reconstructed four-vector components of the H -> bb candidate in the resolved topology."),
    ("HH system", "`HH_pt`, `HH_eta`, `HH_phi`, `HH_m`", "Reconstructed Higgs-pair system in the resolved topology."),
    ("Boosted HH system", "`boosted_HH_pt`, `boosted_HH_eta`, `boosted_HH_phi`, `boosted_HH_m`", "Reconstructed Higgs-pair system in the boosted topology."),
    ("Alternative object labels", "`H1_*`, `H2_*`, `b1_*`, `b2_*`, `g1_*`, `g2_*`", "Auxiliary reconstructed Higgs, b-jet, and photon object four-vectors kept for analysis cross-checks."),
    ("Voxel inputs", "`voxel_photon`, `voxel_jet`, `voxel_bjet`, `voxel_track`, `voxel_tower`", "Calorimeter- or object-image style inputs stored in the original ROOT ntuples and used for 25 x 25 voxel-based ML studies."),
]

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.style = "Table Grid"
headers = ["Variable group", "Variables", "Definition"]
for idx, text in enumerate(headers):
    set_cell_shading(table.rows[0].cells[idx], "F2F4F7")
    write_cell(table.rows[0].cells[idx], text, bold=True)
set_repeat_table_header(table.rows[0])

for group, variables, definition in rows:
    cells = table.add_row().cells
    write_cell(cells[0], group, bold=True)
    write_cell(cells[1], variables.replace("`", ""))
    write_cell(cells[2], definition)

set_table_geometry(table, [1800, 3600, 3960])

doc.add_paragraph()
add_body_paragraph(
    doc,
    "Note: the HDF5 exports contain the event-level variables. The voxel branches are available in the original ROOT ntuples used to produce the ML inputs.",
)

doc.save(OUT)
print(OUT)
